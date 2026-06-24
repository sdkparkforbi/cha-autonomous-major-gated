# -*- coding: utf-8 -*-
"""블록 → DOCX (python-docx). index.html을 최대한 반영."""
import re, pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from parse_proposal import parse

HERE = pathlib.Path(__file__).parent
NAVY = RGBColor(0x15, 0x29, 0x4D); NAVY2 = RGBColor(0x1F, 0x3A, 0x6E)
TEAL = RGBColor(0x0E, 0x8A, 0x7D); ROSE = RGBColor(0xB2, 0x3A, 0x48)
AMBER = RGBColor(0x8A, 0x66, 0x00); INK = RGBColor(0x1C, 0x22, 0x30); MUT = RGBColor(0x5B, 0x64, 0x77)
BOXFILL = {"key": "EEF3FB", "hyp": "E7F5F2", "gap": "FBE9EB", "dropin": "FBF3DF"}
BOXCLR = {"key": NAVY2, "hyp": TEAL, "gap": ROSE, "dropin": AMBER}
FONT = "맑은 고딕"

def set_kfont(run):
    run.font.name = FONT
    r = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), FONT); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT)

def shade(el, hexcolor):
    tcPr = el.get_or_add_tcPr() if hasattr(el, "get_or_add_tcPr") else el
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def cell_shade(cell, hexcolor):
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)

def add_runs(p, runs, size=10.5, base_color=INK, force_bold=False):
    for st, t in runs:
        if t == "\n":
            p.add_run().add_break(); continue
        if not t:
            continue
        r = p.add_run(t); r.font.size = Pt(size); r.font.color.rgb = base_color
        r.bold = force_bold or (st in ("b", "em"))
        if st == "em":
            r.font.color.rgb = NAVY2
        set_kfont(r)
    return p

NUMRE = re.compile(r"^[\d,]+(\.\d+)?\s*(천원|%)?$")

def add_table(doc, tb):
    if tb["caption"]:
        cap = doc.add_paragraph(); cap.paragraph_format.space_before = Pt(6); cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(tb["caption"]); r.font.size = Pt(9); r.font.color.rgb = MUT; r.bold = True; set_kfont(r)
    ncols = tb["ncols"]
    t = doc.add_table(rows=0, cols=ncols); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    has_head = bool(tb["headers"])
    if has_head:
        row = t.add_row().cells
        for i, h in enumerate(tb["headers"]):
            cell_shade(t.rows[0].cells[i], "15294D")
            p = row[i].paragraphs[0]; rr = p.add_run(h); rr.bold = True; rr.font.size = Pt(9.5)
            rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); set_kfont(rr)
    for ri, cells in enumerate(tb["rows"]):
        row = t.add_row()
        rcells = row.cells
        is_total = any(("소계" in txt or "합계" in txt) for txt, _ in cells)
        col = 0
        for (txt, cs) in cells:
            target = rcells[col]
            if cs > 1:
                for k in range(1, cs):
                    target = target.merge(rcells[col + k])
            p = target.paragraphs[0]
            if NUMRE.match(txt):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rr = p.add_run(txt); rr.font.size = Pt(9.5); rr.bold = is_total; set_kfont(rr)
            if is_total:
                if "합계" in (cells[0][0]):
                    cell_shade(target, "15294D"); rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    cell_shade(target, "EEF3FB")
            col += cs
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_heading(doc, text, level):
    p = doc.add_paragraph(); pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(16); pf.space_after = Pt(6)
        r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
        pb = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18"); bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), "15294D")
        pb.append(bottom); p._p.get_or_add_pPr().append(pb)
    else:
        pf.space_before = Pt(10); pf.space_after = Pt(3)
        r = p.add_run("▍ " + text); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = NAVY2
    set_kfont(r)

def add_box(doc, box):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    cell = t.cell(0, 0); cell_shade(cell, BOXFILL.get(box["kind"], "EEF3FB"))
    p = cell.paragraphs[0]; rr = p.add_run(box["tag"]); rr.bold = True; rr.font.size = Pt(10)
    rr.font.color.rgb = BOXCLR.get(box["kind"], NAVY2); set_kfont(rr)
    inner = []
    for blk in box["blocks"]:
        if blk["type"] == "p":
            pp = cell.add_paragraph(); add_runs(pp, blk["runs"], size=10)
        else:
            inner.append(blk)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    for tbl in inner:
        add_table(doc, tbl)

def build():
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10.5)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    sec = doc.sections[0]
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2); sec.left_margin = Cm(2); sec.right_margin = Cm(2)

    for b in parse():
        ty = b["type"]
        if ty == "title":
            kp = doc.add_paragraph(); kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            kr = kp.add_run(b["kicker"]); kr.font.size = Pt(10); kr.font.color.rgb = TEAL; kr.bold = True; set_kfont(kr)
            tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr = tp.add_run(b["title"]); tr.font.size = Pt(26); tr.bold = True; tr.font.color.rgb = NAVY; set_kfont(tr)
            sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(b["sub"]); sr.font.size = Pt(13); sr.font.color.rgb = NAVY2; set_kfont(sr)
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run("  |  ".join(b["chips"])); cr.font.size = Pt(9); cr.font.color.rgb = MUT; set_kfont(cr)
            for spans in b["metas"]:
                mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mr = mp.add_run("   ".join(spans)); mr.font.size = Pt(9.5); mr.font.color.rgb = INK; set_kfont(mr)
            doc.add_paragraph()
        elif ty == "note":
            t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
            cell = t.cell(0, 0); cell_shade(cell, "0E8A7D")
            p = cell.paragraphs[0]; add_runs(p, b["runs"], size=10, base_color=RGBColor(0xFF, 0xFF, 0xFF))
            for r in p.runs: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            doc.add_paragraph()
        elif ty == "h1":
            if b["text"].startswith(("우리의 전략", "(예상) 결과")):
                doc.add_page_break()
            add_heading(doc, b["text"], 1)
        elif ty == "h2":
            add_heading(doc, b["text"], 2)
        elif ty == "p":
            p = doc.add_paragraph(); add_runs(p, b["runs"], size=11 if b.get("lead") else 10.5,
                                              base_color=NAVY2 if b.get("lead") else INK)
        elif ty == "ul":
            for item in b["items"]:
                p = doc.add_paragraph(style="List Bullet"); add_runs(p, item, size=10.5)
        elif ty == "box":
            add_box(doc, b)
        elif ty == "table":
            add_table(doc, b)
        elif ty == "figure":
            p = doc.add_paragraph(); r = p.add_run("[그림] GATED 파이프라인: 비정형 텍스트 → 인과 지식그래프 → 그래프 추론")
            r.bold = True; r.font.color.rgb = NAVY; set_kfont(r)
            p2 = doc.add_paragraph()
            add_runs(p2, [("", "STEP 1 대량 수집(뉴스 API·KCI) → STEP 2 관련성 필터링(GPT) → "
                          "STEP 3 관계 추출·지식그래프 → STEP 4 GATED 추론(허브·예측)")], size=10)
            p3 = doc.add_paragraph()
            add_runs(p3, [("", "추출 triplet 예시: (무전공 유형Ⅰ→촉진→인기학과 쏠림) / (전공탐색학기→억제→중도탈락)")], size=9.5)
            cp = doc.add_paragraph(); cr = cp.add_run(b["caption"]); cr.font.size = Pt(9); cr.font.color.rgb = MUT; set_kfont(cr)
        elif ty == "foot":
            doc.add_paragraph()
            fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = fp.add_run("  ·  ".join(b["spans"])); fr.font.size = Pt(8.5); fr.font.color.rgb = MUT; set_kfont(fr)

    out = HERE / "자율전공_GATED_연구제안서.docx"
    doc.save(str(out))
    print("DOCX:", out, out.stat().st_size, "bytes")

if __name__ == "__main__":
    build()
